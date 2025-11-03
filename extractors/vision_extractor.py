import os
import io
import base64
from typing import List, Dict, Any
from pathlib import Path
import subprocess

try:
    import pypdfium2 as pdfium
    from PIL import Image
except Exception:
    pdfium = None  # type: ignore
    Image = None  # type: ignore

try:
    from openai import OpenAI
except Exception:
    OpenAI = None  # type: ignore

try:
    from docx import Document
except Exception:
    Document = None  # type: ignore

try:
    from docx2pdf import convert
except Exception:
    convert = None  # type: ignore

import tempfile  # Built-in module, should always be available

SYSTEM_PROMPT = (
    "You are an expert resume parser. You will receive one or more images of a resume (from PDF or DOCX files)."
    " Extract accurate, structured JSON with these keys:"
    " ['name','email','phone','linkedin','github','highest_degree','college_name','graduation_year','major','cgpa',"
    " 'total_experience_years','current_company','current_designation','previous_companies','technical_skills',"
    " 'programming_languages','frameworks_tools','soft_skills','certifications']."
    " If not present, use 'Not Found'. Return only JSON."
)

USER_PROMPT = (
    "Parse the attached resume images and return the JSON fields."
)


def _pdf_to_images(pdf_bytes: bytes, max_pages: int = 3, scale: float = 2.0) -> List[Image.Image]:
    if pdfium is None or Image is None:
        return []
    images: List[Image.Image] = []
    try:
        pdf = pdfium.PdfDocument(io.BytesIO(pdf_bytes))
        page_count = len(pdf)
        for i in range(min(page_count, max_pages)):
            page = pdf[i]
            pil_image = page.render(scale=scale).to_pil()
            images.append(pil_image.convert("RGB"))
    except Exception as e:
        print(f"[WARN] PDF rendering failed: {e}")
    return images


def _docx_to_images(docx_bytes: bytes, max_pages: int = 3, scale: float = 2.0) -> List[Image.Image]:
    """
    Convert DOCX file to images by first converting to PDF, then to images.
    Falls back to text extraction if conversion fails.
    """
    
    if pdfium is None:
        print("[WARNING] DOCX to Images: pypdfium2 not available")
    if Image is None:
        print("[WARNING] DOCX to Images: PIL/Pillow not available")
    if Document is None:
        print("[WARNING] DOCX to Images: python-docx not available")
    if convert is None:
        print("[WARNING] DOCX to Images: docx2pdf not available")
    if pdfium is None or Image is None or Document is None:
        print("[ERROR] DOCX to Images: Required libraries not available, trying alternative method...")
        # Try alternative method
        try:
            return _docx_to_images_alternative(docx_bytes, max_pages=max_pages)
        except Exception as e:
            print(f"[ERROR] DOCX to Images: Alternative method also failed: {e}")
            return []
    
    images: List[Image.Image] = []
    tmp_docx_path = None
    tmp_pdf_path = None
    
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
            tmp_docx.write(docx_bytes)
            tmp_docx_path = tmp_docx.name
        print(f"[DEBUG] DOCX to Images: Created temp DOCX file: {tmp_docx_path}")
        
        # Try LibreOffice (soffice) first for Linux-friendly conversion
        tmp_pdf_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
                tmp_pdf_path = tmp_pdf.name
            print("[DEBUG] DOCX to Images: Trying LibreOffice headless conversion...")
            # Write DOCX to a temp directory and request output to same dir
            tmp_dir = os.path.dirname(tmp_pdf_path)
            # LibreOffice requires output dir and input path; it writes file with same basename
            # Ensure input file resides in tmp_dir for predictable output name
            input_path = tmp_docx_path
            cmd = [
                'soffice', '--headless', '--convert-to', 'pdf', '--outdir', tmp_dir, input_path
            ]
            try:
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                # Determine produced PDF path
                produced_pdf = os.path.join(tmp_dir, Path(input_path).with_suffix('.pdf').name)
                if os.path.exists(produced_pdf):
                    with open(produced_pdf, 'rb') as f:
                        pdf_bytes = f.read()
                    images = _pdf_to_images(pdf_bytes, max_pages=max_pages, scale=scale)
                    return images
                else:
                    print("[WARNING] DOCX to Images: LibreOffice did not produce a PDF file")
            except Exception as soffice_e:
                print(f"[WARNING] DOCX to Images: LibreOffice conversion failed: {soffice_e}")
        finally:
            if tmp_pdf_path and os.path.exists(tmp_pdf_path):
                try:
                    os.unlink(tmp_pdf_path)
                except Exception:
                    pass

        # Only use docx2pdf if available (Windows/Mac with Word)
        if convert is not None:
            try:
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
                    tmp_pdf_path = tmp_pdf.name
                
                # Convert DOCX to PDF
                convert(tmp_docx_path, tmp_pdf_path)
                
                # Read PDF and convert to images
                with open(tmp_pdf_path, 'rb') as f:
                    pdf_bytes = f.read()
                
                images = _pdf_to_images(pdf_bytes, max_pages=max_pages, scale=scale)
            except Exception as convert_e:
                print(f"[ERROR] DOCX to Images: docx2pdf conversion error: {convert_e}")
                tmp_pdf_path = None
                raise
        else:
            print("[DEBUG] DOCX to Images: docx2pdf not available, trying alternative method...")
            raise ImportError("docx2pdf not available")
                
    except Exception as e:
        import traceback
        # Fallback: try alternative method using python-docx rendering
        try:
            images = _docx_to_images_alternative(docx_bytes, max_pages=max_pages)
        except Exception as e2:
            import traceback
    finally:
        # Cleanup temporary files
        if tmp_pdf_path:
            try:
                if os.path.exists(tmp_pdf_path):
                    os.unlink(tmp_pdf_path)
                    print(f"[DEBUG] DOCX to Images: Cleaned up temp PDF: {tmp_pdf_path}")
            except Exception as cleanup_e:
                print(f"[WARNING] DOCX to Images: Failed to cleanup PDF: {cleanup_e}")
        if tmp_docx_path:
            try:
                if os.path.exists(tmp_docx_path):
                    os.unlink(tmp_docx_path)
                    print(f"[DEBUG] DOCX to Images: Cleaned up temp DOCX: {tmp_docx_path}")
            except Exception as cleanup_e:
                print(f"[WARNING] DOCX to Images: Failed to cleanup DOCX: {cleanup_e}")
    
    return images


def _docx_to_images_alternative(docx_bytes: bytes, max_pages: int = 3) -> List[Image.Image]:
    """
    Alternative method: Convert DOCX to images by rendering document structure.
    This is a fallback if docx2pdf is not available or fails.
    """
    if Document is None or Image is None:
        return []
    
    images: List[Image.Image] = []
    
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        
        # Parse DOCX
        doc = Document(io.BytesIO(docx_bytes))
        
        # Create PDF in memory
        pdf_buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Extract and add content from DOCX
        para_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
                para_count += 1
                if para_count > 50:  # Limit content to avoid huge PDFs
                    break
        
        # Build PDF
        pdf_doc.build(story)
        pdf_bytes = pdf_buffer.getvalue()
        pdf_buffer.close()
        
        # Convert PDF to images
        images = _pdf_to_images(pdf_bytes, max_pages=max_pages)
        
    except ImportError:
        # reportlab not available, return empty
        print("[WARN] reportlab not available for DOCX rendering")
        return []
    except Exception as e:
        print(f"[WARN] Alternative DOCX rendering failed: {e}")
        return []
    
    return images


def _img_to_data_url(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{b64}"


def extract_with_openai_vision(file_bytes: bytes, file_type: str = 'pdf') -> Dict[str, Any]:
    """
    Extract resume data using OpenAI Vision API.
    
    Args:
        file_bytes: The file content as bytes (PDF or DOCX)
        file_type: 'pdf' or 'docx' to specify file type
    
    Returns:
        Dictionary with extracted resume fields
    """
    
    if not file_bytes:
        print("[ERROR] Vision Extractor: Empty file bytes")
        return {}
    if OpenAI is None:
        print("[ERROR] Vision Extractor: OpenAI library not available")
        return {}
    client = None
    try:
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            print("[ERROR] Vision Extractor: OPENAI_API_KEY not found in environment")
            return {}
        client = OpenAI(api_key=api_key)
    except Exception as e:
        print(f"[ERROR] Vision Extractor: Failed to initialize OpenAI client: {e}")
        return {}

    # Convert file to images based on type
    if file_type.lower() == 'docx':
        images = _docx_to_images(file_bytes)
    else:
        images = _pdf_to_images(file_bytes)
        
    if not images:
        print("[WARNING] Vision Extractor: No images generated, cannot proceed with vision API")
        return {}

    contents = [{"type": "text", "text": USER_PROMPT}]
    for idx, img in enumerate(images):
        contents.append({
            "type": "image_url",
            "image_url": {"url": _img_to_data_url(img)}
        })

    try:
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": contents}
            ],
            temperature=0.0,
        )
        text = (resp.choices[0].message.content or '').strip()
        
        # Attempt to extract JSON from fenced blocks if present
        if text.startswith("```"):
            text = text.strip('`')
            text = text.replace("json", "", 1).strip()
        
        import json
        data = json.loads(text)
        if isinstance(data, dict):
            return data
        return {}
    except json.JSONDecodeError as e:
        print(f"[ERROR] Vision Extractor: JSON decode error: {e}")
        return {}
    except Exception as e:
        import traceback
        print(f"[ERROR] Vision Extractor: OpenAI vision extraction failed: {e}")
        print(f"[ERROR] Vision Extractor: Traceback:\n{traceback.format_exc()}")
        return {}
