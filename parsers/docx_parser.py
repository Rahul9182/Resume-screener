from io import BytesIO
from docx import Document
import traceback

def extract_text_from_docx(file_content):
    """
    Extracts all text from a DOCX file including paragraphs and tables.

    Args:
        file_content (bytes): The file content (from uploaded document)

    Returns:
        str: Combined extracted text
    """
    text = []
    try:
        
        # Validate file content
        if not file_content or len(file_content) == 0:
            return ""
        
        # Check if it's actually a DOCX file (starts with PK header for ZIP format)
        if not file_content.startswith(b'PK'):
            return ""
        
        document = Document(BytesIO(file_content))
        
        # Extract paragraphs (including those in headers/footers)
        para_count = 0
        for para in document.paragraphs:
            if para.text and para.text.strip():
                text.append(para.text.strip())
                para_count += 1
        
        # Extract from headers
        header_count = 0
        for section in document.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text and para.text.strip():
                        text.append(para.text.strip())
                        header_count += 1
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text and para.text.strip():
                        text.append(para.text.strip())
                        header_count += 1
        if header_count > 0:
            print(f"[DEBUG] DOCX Parser: Extracted {header_count} paragraphs from headers/footers")
        
        # Extract tables (resumes often have tables for skills, experience, etc.)
        table_count = 0
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    # Extract paragraphs from cells too (cells can have multiple paragraphs)
                    cell_paras = []
                    for para in cell.paragraphs:
                        if para.text and para.text.strip():
                            cell_paras.append(para.text.strip())
                    if cell_paras:
                        row_text.append(' '.join(cell_paras))
                if row_text:
                    text.append(' | '.join(row_text))
                    table_count += 1
        
        # Try to extract from text boxes (if any) - this requires accessing the document.xml
        try:
            # Access document body to find text boxes and other inline elements
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.text.run import CT_R
            
            # Additional extraction: try to get all text nodes from XML
            body = document._body._body
            all_texts = []
            for paragraph in body.xpath('.//w:t'):
                if paragraph.text:
                    all_texts.append(paragraph.text.strip())
            
            # Add unique text that wasn't already captured
            existing_text = ' '.join(text).lower()
            new_texts = [t for t in all_texts if t and t.lower() not in existing_text and len(t.strip()) > 1]
            if new_texts:
                text.extend(new_texts)
        except Exception as xml_e:
            print(f"[DEBUG] DOCX Parser: XML-based extraction not available or failed: {xml_e}")
        
        # Combine all text
        combined_text = '\n'.join(text)
        
        if len(combined_text.strip()) < 10:
            print("[WARNING] DOCX Parser: Very little text extracted. Document might be empty or contain only images.")
        
        return combined_text
        
    except Exception as e:
        error_msg = f"[ERROR] DOCX Parser: Exception occurred - {str(e)}"
        print(error_msg)
        print(f"[ERROR] DOCX Parser: Traceback:\n{traceback.format_exc()}")
        return ""